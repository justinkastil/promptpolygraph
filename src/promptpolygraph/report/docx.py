"""DOCX rendering of a polygraph run review via python-docx.

Mirrors the Markdown report's structure (cover, category-score table,
worst-first per-category cases, persona panel, forensic synthesis, A/B, cost &
latency) as a real Word document, returned as raw bytes. Robust to None/missing
optional inputs.
"""

from __future__ import annotations

from io import BytesIO
from typing import Any, Optional

from docx import Document
from docx.shared import Pt, RGBColor

from ..models import Case, Response, Rubric, RunMeta, Score
from .context import _build_persona_comparison, _build_personas
from .markdown import _fmt_delta, _fmt_num, _verdict_rank

_MONO = "Courier New"
_TABLE_STYLE = "Light Grid Accent 1"


def _mark(ok: Any) -> str:
    return "Pass" if ok else "Fail"


def _try_table_style(table, doc: Document) -> None:
    try:
        table.style = _TABLE_STYLE
    except Exception:  # noqa: BLE001 - style may not exist in the template
        try:
            table.style = "Table Grid"
        except Exception:  # noqa: BLE001
            pass


def _mono_block(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text if text else "(empty)")
    run.font.name = _MONO
    run.font.size = Pt(9)


def _kv(doc: Document, key: str, value: Any) -> None:
    p = doc.add_paragraph()
    r = p.add_run(f"{key}: ")
    r.bold = True
    p.add_run("—" if value in (None, "") else str(value))


def _cover(doc: Document, run_meta: RunMeta, summary: dict, cases: list[Case], scores: list[Score]) -> None:
    doc.add_heading(f"Polygraph Review — {run_meta.name}", level=0)
    _kv(doc, "Run ID", run_meta.run_id)
    _kv(doc, "Adapter", run_meta.adapter)
    _kv(doc, "Model", run_meta.model)
    _kv(doc, "Mode", run_meta.mode)
    _kv(doc, "Created", run_meta.created_at)
    _kv(doc, "Completed", run_meta.completed_at)
    _kv(doc, "Cases executed", f"{run_meta.completed_cases or len(cases)} / {run_meta.total_cases or len(cases)}")
    _kv(doc, "Cases analyzed", len(scores))
    _kv(doc, "Threshold", _fmt_num(summary.get("threshold"), 1))

    passing = summary.get("categories_passing", 0)
    total = summary.get("categories_total", 0)
    overall = summary.get("overall_pass", False)
    h = doc.add_heading(level=1)
    run = h.add_run(f"Overall verdict: {'PASS' if overall else 'FAIL'} — {passing}/{total} categories passing")
    try:
        run.font.color.rgb = RGBColor(0x1B, 0x7F, 0x37) if overall else RGBColor(0xB0, 0x1B, 0x1B)
    except Exception:  # noqa: BLE001
        pass


def _category_table(doc: Document, summary: dict, baseline_diff: Optional[dict]) -> None:
    doc.add_heading("Category scores", level=1)
    dims = summary.get("dimensions") or []
    cat_scores = summary.get("category_scores") or {}
    has_base = bool(baseline_diff and baseline_diff.get("by_category"))

    cols = ["Category", "Count", *dims, "Pass"] + (["Δ vs baseline"] if has_base else [])
    table = doc.add_table(rows=1, cols=len(cols))
    _try_table_style(table, doc)
    hdr = table.rows[0].cells
    for i, name in enumerate(cols):
        hdr[i].text = name

    for cat in sorted(cat_scores):
        entry = cat_scores[cat] or {}
        cells = table.add_row().cells
        cells[0].text = cat
        cells[1].text = str(entry.get("count", 0))
        idx = 2
        for d in dims:
            cells[idx].text = _fmt_num(entry.get(d))
            idx += 1
        cells[idx].text = _mark(entry.get("pass"))
        idx += 1
        if has_base:
            deltas = []
            base_tab = (baseline_diff.get("by_category") or {}).get(cat) or {}
            for d in dims:
                e = base_tab.get(d)
                if isinstance(e, dict) and e.get("delta") is not None:
                    deltas.append(f"{d} {_fmt_delta(e.get('delta'))}")
            cells[idx].text = ", ".join(deltas) if deltas else "—"


def _case_detail(doc: Document, case: Case, resp: Optional[Response], score: Optional[Score], dims: list[str]) -> None:
    verdict = "PASS" if (score and score.verdict_pass) else "FAIL"
    doc.add_heading(f"Case {case.id} — {verdict}", level=3)
    _kv(doc, "Prompt", case.prompt)
    if case.expected_behavior:
        _kv(doc, "Expected behavior", case.expected_behavior)
    if case.red_flags:
        p = doc.add_paragraph()
        p.add_run("Red flags:").bold = True
        for rf in case.red_flags:
            doc.add_paragraph(str(rf), style="List Bullet")
    doc.add_paragraph().add_run("Response:").bold = True
    _mono_block(doc, (resp.text if resp and resp.text else "") + (f"\n[error: {resp.error}]" if resp and resp.error else ""))
    if score:
        bits = []
        for d in dims:
            v = score.dimensions.get(d)
            bits.append(f"{d}={_fmt_num(v, 0) if v is not None else 'n/a'}")
        _kv(doc, "Scores", ", ".join(bits))
        if score.assertions:
            p = doc.add_paragraph()
            p.add_run("Assertions:").bold = True
            for a in score.assertions:
                doc.add_paragraph(
                    f"[{_mark(a.passed)}] {a.kind} {a.description or ''} {a.detail or ''}".rstrip(),
                    style="List Bullet",
                )
        if score.failure_reason:
            _kv(doc, "Failure reason", score.failure_reason)
        if score.notes:
            _kv(doc, "Notes", score.notes)


def _per_category_sections(
    doc: Document,
    cases: list[Case],
    responses: list[Response],
    scores: list[Score],
    summary: dict,
) -> None:
    doc.add_heading("Per-category detail", level=1)
    dims = summary.get("dimensions") or []
    resp_by_id = {r.case_id: r for r in responses}
    score_by_id = {s.case_id: s for s in scores}
    cats: dict[str, list[Case]] = {}
    for c in cases:
        cats.setdefault(c.category or "default", []).append(c)
    for cat in sorted(cats):
        doc.add_heading(cat, level=2)
        ordered = sorted(cats[cat], key=lambda c: _verdict_rank(score_by_id.get(c.id)))
        for c in ordered:
            _case_detail(doc, c, resp_by_id.get(c.id), score_by_id.get(c.id), dims)


def _persona_panel(doc: Document, audit: Optional[dict]) -> None:
    personas = _build_personas(audit)
    comparison = _build_persona_comparison(audit)
    if not personas and not comparison:
        return
    doc.add_heading("Persona panel", level=1)
    for p in personas:
        doc.add_heading(str(p.get("persona") or "Persona"), level=2)
        summ = p.get("summary") or ""
        if summ:
            doc.add_paragraph(str(summ))
        bits = []
        if p.get("avg_trust") is not None:
            bits.append(f"trust {p['avg_trust']}")
        if p.get("avg_usefulness") is not None:
            bits.append(f"usefulness {p['avg_usefulness']}")
        if p.get("avg_clarity") is not None:
            bits.append(f"clarity {p['avg_clarity']}")
        if p.get("would_return_rate") is not None:
            bits.append(f"would-return {p['would_return_rate']}%")
        if bits:
            _kv(doc, "Scores", " · ".join(bits))
        if p.get("verdict_counts"):
            _kv(doc, "Verdicts", ", ".join(f"{v} ({n})" for v, n in p["verdict_counts"].items()))
        frustrations = p.get("biggest_frustrations") or []
        if frustrations:
            doc.add_paragraph().add_run("Biggest frustrations:").bold = True
            for f in frustrations:
                doc.add_paragraph(str(f), style="List Bullet")
    if comparison:
        doc.add_heading("Methodology check", level=2)
        if comparison.get("rubric_fidelity_verdict"):
            _kv(doc, "Rubric fidelity verdict", comparison["rubric_fidelity_verdict"])
        if comparison.get("divergence_count") is not None:
            _kv(doc, "Rubric-vs-persona divergences", comparison["divergence_count"])
        if comparison.get("blind_spots"):
            doc.add_paragraph().add_run("Human-value blind spots:").bold = True
            for b in comparison["blind_spots"]:
                doc.add_paragraph(str(b), style="List Bullet")
        if comparison.get("final_path"):
            _kv(doc, "Final path", comparison["final_path"])


def _forensic_synthesis(doc: Document, audit: Optional[dict]) -> None:
    doc.add_heading("Forensic synthesis", level=1)
    forensic = (audit or {}).get("forensic") or {}
    synthesis = forensic.get("synthesis") or {}
    if not synthesis and not forensic.get("category_audits"):
        doc.add_paragraph("No forensic synthesis available for this run.")
        return
    patterns = synthesis.get("cross_category_patterns") or synthesis.get("patterns") or []
    if patterns:
        doc.add_heading("Cross-category patterns", level=2)
        for p in patterns:
            txt = (p.get("pattern") or p.get("summary") or str(p)) if isinstance(p, dict) else str(p)
            doc.add_paragraph(txt, style="List Bullet")
    changes = synthesis.get("prioritized_changes") or synthesis.get("ranked_changes") or []
    if changes:
        doc.add_heading("Prioritized changes", level=2)
        for ch in changes:
            if isinstance(ch, dict):
                title = ch.get("change") or ch.get("title") or ch.get("summary") or str(ch)
                why = ch.get("rationale") or ch.get("why") or ""
                txt = f"{title}" + (f" — {why}" if why else "")
            else:
                txt = str(ch)
            doc.add_paragraph(txt, style="List Number")


def _ab_section(doc: Document, pairwise: Optional[dict]) -> None:
    if not pairwise:
        return
    doc.add_heading("A/B comparison", level=1)
    a = pairwise.get("run_a", "a")
    b = pairwise.get("run_b", "b")
    doc.add_paragraph(
        f"{a} wins: {pairwise.get('wins_a', 0)}  ·  {b} wins: {pairwise.get('wins_b', 0)}  ·  ties: {pairwise.get('ties', 0)}"
    )
    by_cat = pairwise.get("by_category") or {}
    if by_cat:
        table = doc.add_table(rows=1, cols=4)
        _try_table_style(table, doc)
        hdr = table.rows[0].cells
        for i, name in enumerate(["Category", a, b, "tie"]):
            hdr[i].text = name
        for cat in sorted(by_cat):
            rec = by_cat[cat] or {}
            cells = table.add_row().cells
            cells[0].text = cat
            cells[1].text = str(rec.get("a", 0))
            cells[2].text = str(rec.get("b", 0))
            cells[3].text = str(rec.get("tie", 0))


def _cost_latency_footer(doc: Document, summary: dict) -> None:
    doc.add_heading("Cost & latency", level=1)
    cost = summary.get("cost") or {}
    lat = summary.get("latency") or {}
    agreement = summary.get("agreement_mean")
    _kv(doc, "Tokens in", f"{cost.get('tokens_in', 0):,}")
    _kv(doc, "Tokens out", f"{cost.get('tokens_out', 0):,}")
    _kv(doc, "Cost (USD)", ("$" + _fmt_num(cost.get("usd"), 4)) if cost.get("usd") is not None else "—")
    _kv(doc, "Latency p50 (ms)", _fmt_num(lat.get("p50_ms"), 1))
    _kv(doc, "Latency p95 (ms)", _fmt_num(lat.get("p95_ms"), 1))
    _kv(doc, "Latency mean (ms)", _fmt_num(lat.get("mean_ms"), 1))
    _kv(doc, "Assertion pass rate", f"{_fmt_num((summary.get('assertion_pass_rate') or 0) * 100, 1)}%")
    _kv(doc, "Judge agreement (mean)", _fmt_num(agreement) if agreement is not None else "—")


def render_docx(
    run_meta: RunMeta,
    cases: list[Case],
    responses: list[Response],
    scores: list[Score],
    summary: dict,
    *,
    rubric: Rubric,
    audit: dict | None = None,
    baseline_diff: dict | None = None,
    pairwise: dict | None = None,
) -> bytes:
    """Render the review as a .docx and return its bytes."""
    _ = rubric
    doc = Document()
    _cover(doc, run_meta, summary, cases, scores)
    _category_table(doc, summary, baseline_diff)
    _per_category_sections(doc, cases, responses, scores, summary)
    _persona_panel(doc, audit)
    _forensic_synthesis(doc, audit)
    _ab_section(doc, pairwise)
    _cost_latency_footer(doc, summary)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
