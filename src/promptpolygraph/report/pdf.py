"""PDF rendering of the review.

Strategy, in order of preference:
  1. Convert the already-rendered .docx via a headless LibreOffice
     (`soffice`/`libreoffice` on PATH) — best fidelity.
  2. Fall back to a plain-text PDF via `reportlab` if it is importable.
  3. If neither is available, return None so the caller can log "pdf skipped".

This function NEVER raises: every failure path collapses to a None return.
"""

from __future__ import annotations

import os
import shutil
import subprocess
import tempfile
from typing import Optional


def _find_soffice() -> Optional[str]:
    for name in ("soffice", "libreoffice"):
        path = shutil.which(name)
        if path:
            return path
    return None


def _convert_with_soffice(docx_bytes: bytes, out_path: str, soffice: str) -> Optional[str]:
    try:
        with tempfile.TemporaryDirectory() as tmp:
            src = os.path.join(tmp, "report.docx")
            with open(src, "wb") as fh:
                fh.write(docx_bytes)
            proc = subprocess.run(
                [soffice, "--headless", "--convert-to", "pdf", "--outdir", tmp, src],
                capture_output=True,
                timeout=180,
            )
            produced = os.path.join(tmp, "report.pdf")
            if proc.returncode != 0 or not os.path.exists(produced):
                return None
            os.makedirs(os.path.dirname(os.path.abspath(out_path)) or ".", exist_ok=True)
            shutil.move(produced, out_path)
            return out_path
    except Exception:  # noqa: BLE001 - never raise
        return None


def _convert_with_reportlab(docx_bytes: bytes, out_path: str) -> Optional[str]:
    try:
        from io import BytesIO

        from docx import Document  # type: ignore
        from reportlab.lib.pagesizes import letter  # type: ignore
        from reportlab.lib.units import inch  # type: ignore
        from reportlab.pdfgen import canvas  # type: ignore
    except Exception:  # noqa: BLE001
        return None

    try:
        doc = Document(BytesIO(docx_bytes))
        lines: list[str] = []
        for para in doc.paragraphs:
            lines.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                lines.append(" | ".join(c.text for c in row.cells))

        os.makedirs(os.path.dirname(os.path.abspath(out_path)) or ".", exist_ok=True)
        c = canvas.Canvas(out_path, pagesize=letter)
        width, height = letter
        margin = 0.75 * inch
        y = height - margin
        c.setFont("Helvetica", 9)
        max_chars = 110
        for raw in lines:
            for chunk_start in range(0, max(len(raw), 1), max_chars):
                chunk = raw[chunk_start:chunk_start + max_chars]
                if y < margin:
                    c.showPage()
                    c.setFont("Helvetica", 9)
                    y = height - margin
                c.drawString(margin, y, chunk)
                y -= 12
        c.save()
        return out_path
    except Exception:  # noqa: BLE001
        return None


def render_pdf(docx_bytes: bytes, out_path: str) -> str | None:
    """Convert docx bytes to a PDF at `out_path`. Returns the path or None."""
    if not docx_bytes:
        return None
    soffice = _find_soffice()
    if soffice:
        result = _convert_with_soffice(docx_bytes, out_path, soffice)
        if result:
            return result
    return _convert_with_reportlab(docx_bytes, out_path)
