from __future__ import annotations

import io

from docx import Document

from promptpolygraph import analyze as A
from promptpolygraph.models import Case, Response, RunMeta
from promptpolygraph.report import render_docx, render_html, render_markdown


def _fixture():
    cases = [
        Case(prompt="q1", category="accuracy"),
        Case(prompt="q2", category="tone"),
    ]
    responses = [Response(case_id=c.id, text=f"answer to {c.prompt}", latency_ms=100) for c in cases]
    rubric = A.default_rubric()
    import asyncio

    scores = asyncio.run(A.analyze_run(cases, responses, rubric, mock=True))
    summary = A.summarize(cases, responses, scores, rubric)
    meta = RunMeta(name="t", adapter="demo")
    return meta, cases, responses, scores, summary, rubric


def _audit_fixture():
    return {
        "forensic": {},
        "persona": {
            "reactions": [
                {
                    "persona": "Skeptical Retiree",
                    "persona_summary": "Wants plain answers and distrusts jargon.",
                    "biggest_frustrations": [
                        "Too many hedging caveats",
                        "Assumes I know technical terms",
                    ],
                    "what_would_win_me": ["A direct answer first", "Sources I can check"],
                    "reactions": [
                        {
                            "case_id": "c1",
                            "category": "accuracy",
                            "trust": 6,
                            "usefulness": 7,
                            "clarity": 8,
                            "would_return": True,
                            "verdict": "lukewarm",
                        },
                        {
                            "case_id": "c2",
                            "category": "tone",
                            "trust": 4,
                            "usefulness": 5,
                            "clarity": 6,
                            "would_return": False,
                            "verdict": "would not return",
                        },
                    ],
                },
                {
                    "persona": "Busy Clinician",
                    "persona_summary": "Needs fast, scannable, defensible output.",
                    "biggest_frustrations": ["Walls of text waste my time"],
                    "what_would_win_me": "Bullet-point summaries up top.",
                    "reactions": [
                        {
                            "case_id": "c1",
                            "category": "accuracy",
                            "trust": 9,
                            "usefulness": 9,
                            "clarity": 7,
                            "would_return": "yes",
                            "verdict": "delighted",
                        },
                        {
                            "case_id": "c2",
                            "category": "tone",
                            "trust": 8,
                            "usefulness": 8,
                            "clarity": 9,
                            "would_return": "yes",
                            "verdict": "delighted",
                        },
                    ],
                },
            ],
            "comparison": {
                "divergences": [
                    {"case_id": "c2", "type": "tone", "evidence": "rubric passed, humans bounced"},
                    {"case_id": "c1", "type": "trust", "evidence": "humans split on trust"},
                ],
                "rubric_fidelity_verdict": "Rubric mostly tracks human judgment, with tone gaps.",
                "chasing_tail_risks": ["Over-optimizing adversarial edge cases"],
                "human_value_blindspots": [
                    "Conciseness is unscored but matters",
                    "Source transparency drives trust",
                ],
                "reconciled_priorities": ["Lead with the answer", "Add a sources line"],
                "final_path": "Ship the concise-answer change first.",
            },
        },
    }


def test_markdown_persona_section_present():
    meta, cases, responses, scores, summary, rubric = _fixture()
    audit = _audit_fixture()
    md = render_markdown(meta, cases, responses, scores, summary, rubric=rubric, audit=audit)
    assert "Persona" in md
    assert "Wants plain answers and distrusts jargon." in md  # a persona summary
    assert "Too many hedging caveats" in md  # a frustration
    assert "Methodology check" in md
    assert "Rubric mostly tracks human judgment" in md  # rubric-fidelity verdict
    assert "Conciseness is unscored but matters" in md  # blind spot


def test_html_persona_section_and_radar():
    meta, cases, responses, scores, summary, rubric = _fixture()
    audit = _audit_fixture()
    html = render_html(meta, cases, responses, scores, summary, rubric=rubric, audit=audit)
    assert "Needs fast, scannable, defensible output." in html  # persona summary text
    assert "<svg" in html  # existing radar still rendered
    assert "Methodology check" in html
    assert "Rubric mostly tracks human judgment" in html  # methodology-check verdict
    assert "Would return" in html  # stat chip label
    assert "Human-value blind spots" in html


def test_html_escapes_script_in_persona_text():
    meta, cases, responses, scores, summary, rubric = _fixture()
    audit = _audit_fixture()
    audit["persona"]["reactions"][0]["persona_summary"] = "<script>alert(1)</script>"
    html = render_html(meta, cases, responses, scores, summary, rubric=rubric, audit=audit)
    assert "<script>alert(1)</script>" not in html
    assert "&lt;script&gt;alert(1)&lt;/script&gt;" in html


def test_docx_persona_panel_heading():
    meta, cases, responses, scores, summary, rubric = _fixture()
    audit = _audit_fixture()
    data = render_docx(meta, cases, responses, scores, summary, rubric=rubric, audit=audit)
    assert isinstance(data, bytes) and data
    doc = Document(io.BytesIO(data))
    headings = [p.text for p in doc.paragraphs if p.style and p.style.name.startswith("Heading")]
    assert "Persona panel" in headings
    assert "Methodology check" in headings
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Skeptical Retiree" in full_text
    assert "Too many hedging caveats" in full_text


def test_renders_without_persona_data():
    meta, cases, responses, scores, summary, rubric = _fixture()
    # audit=None must render every format with no persona section and no error.
    md = render_markdown(meta, cases, responses, scores, summary, rubric=rubric, audit=None)
    assert "No persona reactions available" in md
    html = render_html(meta, cases, responses, scores, summary, rubric=rubric, audit=None)
    assert "No persona reactions available" in html
    data = render_docx(meta, cases, responses, scores, summary, rubric=rubric, audit=None)
    doc = Document(io.BytesIO(data))
    headings = [p.text for p in doc.paragraphs if p.style and p.style.name.startswith("Heading")]
    assert "Persona panel" not in headings
